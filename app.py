"""
DocExtract AI - Flask Backend (Groq)
Real OCR: extracts actual text from PDF/DOCX before sending to AI.
"""

from dotenv import load_dotenv
load_dotenv()

import os, uuid, json, base64, mimetypes
from datetime import datetime
from pathlib import Path
from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
from groq import Groq

app = Flask(__name__)
CORS(app, origins=["*"])

UPLOAD_FOLDER = Path("uploads")
STORAGE_FILE  = Path("storage/documents.json")
UPLOAD_FOLDER.mkdir(exist_ok=True)
STORAGE_FILE.parent.mkdir(exist_ok=True)

ALLOWED_EXTS  = {".pdf", ".jpg", ".jpeg", ".png", ".docx", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024

client = Groq(api_key=os.environ.get("GROQ_API_KEY", ""))

# ── Storage ───────────────────────────────────────────────────────────────────

def load_documents():
    if STORAGE_FILE.exists():
        with open(STORAGE_FILE) as f:
            return json.load(f)
    return {}

def save_documents(docs):
    with open(STORAGE_FILE, "w") as f:
        json.dump(docs, f, indent=2)

def now_str():
    return datetime.now().strftime("%b %d, %Y %H:%M")

def allowed_file(filename):
    return Path(filename).suffix.lower() in ALLOWED_EXTS

def file_to_base64(path):
    mime, _ = mimetypes.guess_type(str(path))
    mime = mime or "application/octet-stream"
    with open(path, "rb") as f:
        return base64.standard_b64encode(f.read()).decode(), mime

# ── Real Text Extraction ──────────────────────────────────────────────────────

def extract_text_from_file(file_path: Path, ext: str) -> str:
    """Extract actual text content from PDF, DOCX, or TXT files."""
    text = ""

    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    # Also try extracting tables
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            for row in table:
                                row_text = " | ".join(str(cell) for cell in row if cell)
                                if row_text.strip():
                                    page_text += "\n" + row_text
                    if page_text.strip():
                        pages_text.append(f"[Page {i+1}]\n{page_text.strip()}")
                text = "\n\n".join(pages_text)
            print(f"[OCR] PDF extracted: {len(text)} characters from {len(pdf.pages)} pages")
        except Exception as e:
            print(f"[OCR] pdfplumber failed: {e}")

    elif ext == ".docx":
        try:
            import docx as docxlib
            doc = docxlib.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            text = "\n".join(paragraphs)
            print(f"[OCR] DOCX extracted: {len(text)} characters")
        except Exception as e:
            print(f"[OCR] docx failed: {e}")

    elif ext == ".txt":
        try:
            with open(file_path, "r", errors="ignore") as f:
                text = f.read()
            print(f"[OCR] TXT read: {len(text)} characters")
        except Exception as e:
            print(f"[OCR] txt failed: {e}")

    return text.strip()


# ── Prompt Builder ────────────────────────────────────────────────────────────

def build_prompt(doc_type: str, expected_fields: list, has_real_text: bool) -> str:
    is_auto = doc_type.lower() in ("auto detect", "auto", "")

    if is_auto:
        type_instruction = """First, identify what type of document this is based on its content.
Choose the most accurate type: Invoice, Receipt, Contract, Identity Document, Resume, Medical Report,
Bank Statement, Academic Certificate, Insurance Policy, Payslip, Purchase Order, Tax Document,
Legal Document, Delivery Note, or describe it accurately if none fit."""
        fields_instruction = "Extract ALL fields you can find in the document. Do not skip any field."
    else:
        fields_list = "\n".join(f"- {f}" for f in expected_fields) if expected_fields else f"All relevant fields for a {doc_type}"
        type_instruction = f'This document has been identified as: "{doc_type}". Use this as the document_type.'
        fields_instruction = f"""Extract these specific fields (and any additional relevant fields you find):
{fields_list}

If a field is not present in the document, include it with an empty value and set manual_check to true."""

    source_note = "The document text has been extracted via OCR and is provided above." if has_real_text else "Note: Direct text extraction was not possible. Do your best based on available information."

    return f"""You are an expert document data extraction AI.

{type_instruction}

{fields_instruction}

{source_note}

CRITICAL: Extract the ACTUAL values from the document text provided. Do NOT make up or guess values.
Only include values that are explicitly present in the document text.
If a value is not found, use empty string "" and set manual_check to true.

Return ONLY this JSON (no markdown, no extra text):
{{
  "document_type": "<document type>",
  "overall_confidence": <float 0-100>,
  "ai_insight": "<what this document is, quality of extraction, anything notable>",
  "fields": [
    {{
      "id": "<sequential number as string>",
      "label": "<field name>",
      "value": "<ACTUAL extracted value from document, or empty string>",
      "confidence": <float 0-100>,
      "icon": "<one of: storefront, calendar_today, payments, account_balance, person, description, badge, location_on>",
      "manual_check": <true if value missing or unclear, else false>
    }}
  ]
}}

Rules:
- Values must come ONLY from the actual document text
- Never invent or assume values not present in the text
- overall_confidence reflects how much real data was extracted
- Return ONLY valid JSON"""


# ── AI Extraction ─────────────────────────────────────────────────────────────

def extract_with_ai(doc_id: str, file_path: Path, filename: str, doc_type: str = "Auto Detect", expected_fields: list = []):
    docs = load_documents()
    docs[doc_id]["status"] = "Processing"
    save_documents(docs)

    try:
        ext = Path(filename).suffix.lower()
        prompt = build_prompt(doc_type, expected_fields, False)

        print(f"[EXTRACT] File: {filename} | Type: {doc_type} | Ext: {ext}")
        key = os.environ.get("GROQ_API_KEY", "")
        print(f"[EXTRACT] API Key: {'YES - ' + key[:8] + '...' if key else 'NO - MISSING!'}")

        if ext in (".jpg", ".jpeg", ".png"):
            # Images: send directly to vision model
            print("[EXTRACT] Vision model (image)...")
            b64, mime = file_to_base64(file_path)
            prompt = build_prompt(doc_type, expected_fields, True)
            response = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                max_tokens=4096,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                        {"type": "text", "text": prompt}
                    ]
                }]
            )

        else:
            # PDF/DOCX/TXT: extract real text first
            print(f"[EXTRACT] Extracting real text from {ext}...")
            real_text = extract_text_from_file(file_path, ext)

            if real_text and len(real_text) > 50:
                print(f"[EXTRACT] Got {len(real_text)} chars of real text — sending to AI...")
                prompt = build_prompt(doc_type, expected_fields, True)
                user_msg = f"""Filename: "{filename}"
Document Type Selected: {doc_type}

--- DOCUMENT CONTENT (extracted via OCR) ---
{real_text[:12000]}
--- END OF DOCUMENT ---

{prompt}"""
            else:
                print(f"[EXTRACT] No text extracted — PDF may be scanned/image-based")
                prompt = build_prompt(doc_type, expected_fields, False)
                user_msg = f"""Filename: "{filename}"
Document Type Selected: {doc_type}

Note: This appears to be a scanned/image-based PDF. Text extraction was not possible.
Please extract what fields you can identify from the document type and filename context,
and mark all fields as manual_check: true since the actual values could not be read.

{prompt}"""

            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                max_tokens=4096,
                messages=[
                    {"role": "system", "content": "You are a document extraction AI. Extract ONLY values actually present in the document text. Return only valid JSON."},
                    {"role": "user", "content": user_msg}
                ]
            )

        raw = response.choices[0].message.content.strip()
        print(f"[EXTRACT] Response: {len(raw)} chars")

        # Strip markdown fences
        if raw.startswith("```"):
            lines = raw.split("\n")
            raw = "\n".join(lines[1:])
        if raw.endswith("```"):
            raw = "\n".join(raw.split("\n")[:-1])
        raw = raw.strip()

        result = json.loads(raw)

        detected_type = result.get("document_type", doc_type)
        confidence    = float(result.get("overall_confidence", 90))
        has_manual    = any(f.get("manual_check") for f in result.get("fields", []))
        status        = "Review Required" if (has_manual or confidence < 80) else "Success"

        docs = load_documents()
        docs[doc_id].update({
            "status":     status,
            "type":       detected_type,
            "confidence": round(confidence, 1),
            "extraction": result,
        })
        save_documents(docs)
        print(f"[EXTRACT] SUCCESS → {detected_type} | {status} | {confidence}%")
        return result

    except Exception as e:
        print(f"[EXTRACT ERROR] {type(e).__name__}: {e}")
        docs = load_documents()
        docs[doc_id]["status"] = "Error"
        docs[doc_id]["error"]  = str(e)
        save_documents(docs)
        raise


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/api/health")
def health():
    return jsonify({"status": "ok", "timestamp": now_str()})


@app.route("/api/documents/upload", methods=["POST"])
def upload_document():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    if not file.filename or not allowed_file(file.filename):
        return jsonify({"error": "Invalid or unsupported file"}), 400
    file.seek(0, 2)
    if file.tell() > MAX_FILE_SIZE:
        return jsonify({"error": "File exceeds 10MB"}), 413
    file.seek(0)

    doc_type = request.form.get("doc_type", "Auto Detect").strip()
    try:
        expected_fields = json.loads(request.form.get("expected_fields", "[]"))
    except Exception:
        expected_fields = []

    doc_id    = str(uuid.uuid4())
    ext       = Path(file.filename).suffix.lower()
    file_path = UPLOAD_FOLDER / f"{doc_id}{ext}"
    file.save(file_path)

    record = {
        "id":         doc_id,
        "name":       file.filename,
        "type":       doc_type if doc_type != "Auto Detect" else "Detecting...",
        "status":     "Processing",
        "date":       now_str(),
        "confidence": 0,
        "file_path":  str(file_path),
        "extraction": None,
        "approved":   False,
    }
    docs = load_documents()
    docs[doc_id] = record
    save_documents(docs)

    try:
        extract_with_ai(doc_id, file_path, file.filename, doc_type, expected_fields)
    except Exception as e:
        print(f"[UPLOAD ERROR] {e}")

    docs = load_documents()
    return jsonify(_public(docs[doc_id])), 201


@app.route("/api/documents", methods=["GET"])
def list_documents():
    docs = list(load_documents().values())
    docs.sort(key=lambda d: d["date"], reverse=True)
    status = request.args.get("status")
    dtype  = request.args.get("type")
    search = request.args.get("search", "").lower()
    if status: docs = [d for d in docs if d["status"] == status]
    if dtype:  docs = [d for d in docs if d["type"] == dtype]
    if search: docs = [d for d in docs if search in d["name"].lower()]
    total    = len(docs)
    page     = max(1, int(request.args.get("page", 1)))
    per_page = min(100, max(1, int(request.args.get("per_page", 10))))
    start    = (page - 1) * per_page
    return jsonify({
        "documents":   [_public(d) for d in docs[start:start + per_page]],
        "total": total, "page": page, "per_page": per_page,
        "total_pages": (total + per_page - 1) // per_page,
    })


@app.route("/api/documents/<doc_id>", methods=["GET"])
def get_document(doc_id):
    docs = load_documents()
    if doc_id not in docs: return jsonify({"error": "Not found"}), 404
    return jsonify(_public_full(docs[doc_id]))


@app.route("/api/documents/<doc_id>/extract", methods=["POST"])
def re_extract(doc_id):
    docs = load_documents()
    if doc_id not in docs: return jsonify({"error": "Not found"}), 404
    fp = Path(docs[doc_id]["file_path"])
    if not fp.exists(): return jsonify({"error": "File missing"}), 410
    try:
        extract_with_ai(doc_id, fp, docs[doc_id]["name"], docs[doc_id].get("type", "Auto Detect"), [])
        docs = load_documents()
        return jsonify({"message": "Done", "document": _public_full(docs[doc_id])})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/documents/<doc_id>/fields", methods=["PATCH"])
def update_fields(doc_id):
    docs = load_documents()
    if doc_id not in docs: return jsonify({"error": "Not found"}), 404
    body    = request.get_json(silent=True) or {}
    updates = {f["id"]: f["value"] for f in body.get("fields", []) if "id" in f}
    doc = docs[doc_id]
    if doc.get("extraction") and "fields" in doc["extraction"]:
        for field in doc["extraction"]["fields"]:
            if field["id"] in updates:
                field["value"]        = updates[field["id"]]
                field["manual_check"] = False
                field["confidence"]   = 100.0
    save_documents(docs)
    return jsonify({"message": "Updated", "document": _public_full(docs[doc_id])})


@app.route("/api/documents/<doc_id>/approve", methods=["POST"])
def approve_document(doc_id):
    docs = load_documents()
    if doc_id not in docs: return jsonify({"error": "Not found"}), 404
    docs[doc_id].update({"approved": True, "status": "Success", "approved_at": now_str()})
    save_documents(docs)
    return jsonify({"message": "Approved", "document": _public(docs[doc_id])})


@app.route("/api/documents/<doc_id>", methods=["DELETE"])
def delete_document(doc_id):
    docs = load_documents()
    if doc_id not in docs: return jsonify({"error": "Not found"}), 404
    doc = docs.pop(doc_id)
    save_documents(docs)
    Path(doc.get("file_path", "")).unlink(missing_ok=True)
    return jsonify({"message": "Deleted"})


@app.route("/api/documents/<doc_id>/file")
def get_file(doc_id):
    docs = load_documents()
    if doc_id not in docs: return jsonify({"error": "Not found"}), 404
    fp = Path(docs[doc_id].get("file_path", ""))
    if not fp.exists(): return jsonify({"error": "File not found"}), 404
    return send_file(fp, download_name=docs[doc_id]["name"])


@app.route("/api/documents/<doc_id>/export")
def export_document(doc_id):
    docs = load_documents()
    if doc_id not in docs: return jsonify({"error": "Not found"}), 404
    doc    = docs[doc_id]
    fmt    = request.args.get("format", "json").lower()
    fields = (doc.get("extraction") or {}).get("fields", [])
    if fmt == "csv":
        import csv, io
        buf = io.StringIO()
        w = csv.DictWriter(buf, fieldnames=["label", "value", "confidence"])
        w.writeheader()
        for f in fields:
            w.writerow({"label": f["label"], "value": f["value"], "confidence": f["confidence"]})
        buf.seek(0)
        return Response(buf.getvalue(), mimetype="text/csv",
                        headers={"Content-Disposition": f"attachment; filename={doc_id}.csv"})
    return jsonify({"document_id": doc_id, "name": doc["name"], "type": doc["type"], "fields": fields})


@app.route("/api/stats")
def get_stats():
    docs  = list(load_documents().values())
    total = len(docs)
    confs = [d["confidence"] for d in docs if d["confidence"] > 0]
    by_type = {}
    for d in docs:
        by_type[d["type"]] = by_type.get(d["type"], 0) + 1
    success = sum(1 for d in docs if d["status"] == "Success")
    return jsonify({
        "total_processed": total,
        "success":         success,
        "failed":          sum(1 for d in docs if d["status"] in ("Error", "Failed")),
        "review_required": sum(1 for d in docs if d["status"] == "Review Required"),
        "processing":      sum(1 for d in docs if d["status"] == "Processing"),
        "success_rate":    round(success / total * 100, 1) if total else 0,
        "avg_confidence":  round(sum(confs) / len(confs), 1) if confs else 0,
        "by_type":         by_type,
    })


def _public(doc):
    return {k: doc[k] for k in ("id", "name", "type", "status", "date", "confidence", "approved") if k in doc}

def _public_full(doc):
    base = _public(doc)
    ext  = doc.get("extraction") or {}
    base.update({
        "fields":      ext.get("fields", []),
        "ai_insight":  ext.get("ai_insight", ""),
        "error":       doc.get("error"),
        "approved_at": doc.get("approved_at"),
    })
    return base


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"🚀 DocExtract AI (Groq) → http://localhost:{port}")
    app.run(host="0.0.0.0", port=port, debug=True)